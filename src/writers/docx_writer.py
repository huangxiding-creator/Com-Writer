"""DOCX 写入器 —— 基于模板生成正式 Word 文档。

策略：复制模板文件 → 替换内容段落 → 保留所有样式。
这样用户的模板格式（字体、字号、间距、编号）被完整保留。
"""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from ..core.models import GeneratedContent
from .template_engine import TemplateInfo, analyze_template
from ..utils.logger import get_logger
from ..config.paths import OUTPUT_DIR

_log = get_logger("writer.docx")


def write_minutes(
    content: GeneratedContent,
    template_info: TemplateInfo,
    template_path: str | Path,
    output_path: str | Path | None = None,
    doc_number: str = "",
) -> str:
    """基于模板生成会议纪要 Word 文档。

    Args:
        content: AI 生成的内容
        template_info: 模板结构信息
        template_path: 模板文件路径
        output_path: 输出路径（None 则自动命名）
        doc_number: 文号（如 〔2026〕2号）
    Returns:
        输出文件路径
    """
    template_path = Path(template_path)

    # 确定输出路径
    if output_path is None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        safe_title = content.title.replace("/", "_").replace("\\", "_")
        output_path = OUTPUT_DIR / f"{safe_title}.docx"
    else:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

    # 复制模板作为基础
    shutil.copy2(template_path, output_path)

    # 加载复制的文档
    doc = Document(str(output_path))

    # 收集要保留的段落索引（标题 + 字段行 + 正文 + 整理人）
    # 策略：找到正文区域，替换正文内容；更新字段值
    paragraphs = doc.paragraphs

    # 标记正文内容段落范围：从 content_start_index 之后到 compiler 之前
    content_start = template_info.content_start_index
    compiler_idx = template_info.compiler_index

    # 构建新段落文本列表（字段值更新 + 正文替换）
    field_updates = _build_field_updates(content, doc_number)
    new_content = list(content.content_paragraphs)

    # 先替换字段值
    for para in paragraphs:
        text = para.text.strip()
        for label, new_value in field_updates.items():
            if label in text:
                _replace_field_value(para, label, new_value)
                break

    # 替换标题
    if content.title and paragraphs:
        # 找到标题段落（第一个有文本的）
        for para in paragraphs:
            if para.text.strip() == template_info.title:
                _replace_paragraph_text(para, content.title)
                break

    # 替换文号
    if doc_number or content.doc_number:
        num = doc_number or content.doc_number
        for para in paragraphs:
            text = para.text.strip()
            if text.startswith("〔") or text.startswith("["):
                _replace_paragraph_text(para, num)
                break

    # 替换正文内容
    _replace_content_paragraphs(
        doc, content_start, compiler_idx, new_content, template_info
    )

    # 替换整理人（未推断出则用模板默认值）
    compiler_name = content.compiler if content.compiler and content.compiler != "待确认" else template_info.compiler
    if compiler_name:
        for para in paragraphs:
            if "整理人" in para.text:
                _replace_field_value(para, "整理人", compiler_name)
                break

    # 保存
    doc.save(str(output_path))

    _log.info("会议纪要已保存: %s", output_path)
    return str(output_path)


def _build_field_updates(content: GeneratedContent, doc_number: str) -> dict[str, str]:
    """构建字段更新映射。"""
    updates: dict[str, str] = {}
    if content.meeting_type:
        updates["会议类别"] = content.meeting_type
    if content.meeting_topic:
        updates["会议议题"] = content.meeting_topic
    if content.meeting_date:
        updates["会议时间"] = content.meeting_date
    if content.meeting_location:
        updates["会议地点"] = content.meeting_location
    if content.host:
        updates["主持人"] = content.host
        updates["主 持 人"] = content.host
    if content.participants:
        updates["参加人员"] = content.participants
    return updates


def _replace_field_value(para: Paragraph, label: str, new_value: str) -> None:
    """替换段落中字段值（保留标签部分）。

    如 "会议类别：专题会" → "会议类别：专题会"（替换冒号后的值）
    """
    text = para.text
    # 找到冒号位置
    for sep in ["：", ":"]:
        if sep in text and label in text:
            prefix = text[: text.find(sep) + 1]
            # 保留段落格式，替换文本
            # 清除现有 runs
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = f"{prefix}{new_value}"
            else:
                para.add_run(f"{prefix}{new_value}")
            return


def _replace_paragraph_text(para: Paragraph, new_text: str) -> None:
    """替换段落全部文本，保留格式。"""
    if para.runs:
        # 保留第一个 run 的格式，清空其余
        first_run = para.runs[0]
        first_run.text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _replace_content_paragraphs(
    doc: Document,
    content_start: int,
    compiler_idx: int,
    new_paragraphs: list[str],
    template_info: TemplateInfo,
) -> None:
    """替换正文内容段落。

    策略：
    1. 找到正文区域（content_start 到 compiler 之间）
    2. 修改第一个正文段落为开头综述
    3. 后续段落按需添加或删除
    """
    if not new_paragraphs:
        return

    all_paras = doc.paragraphs

    # 确定正文区域的段落范围
    # 正文区域：从 content_start+1（"纪要内容："标签之后）到 compiler_idx-1
    body_start = content_start + 1 if content_start > 0 else 1

    # 找到正文段落（跳过空段落）
    body_indices: list[int] = []
    for i in range(body_start, len(all_paras)):
        text = all_paras[i].text.strip()
        if not text:
            continue
        if compiler_idx > 0 and i >= compiler_idx:
            break
        body_indices.append(i)

    if not body_indices:
        _log.warning("未找到正文段落区域，尝试追加")
        # 在"纪要内容："之后添加段落
        for text in new_paragraphs:
            p = doc.add_paragraph(text)
            _apply_style(p, template_info)
        return

    # 获取模板段落的格式作为基准
    template_para = all_paras[body_indices[0]]

    # 替换现有正文段落
    for i, idx in enumerate(body_indices):
        para = all_paras[idx]
        if i < len(new_paragraphs):
            _replace_paragraph_text(para, new_paragraphs[i])
        else:
            # 多余的模板段落，清空
            _replace_paragraph_text(para, "")

    # 如果新段落比模板多，在最后一个正文段落后插入
    if len(new_paragraphs) > len(body_indices):
        last_body_para = all_paras[body_indices[-1]]
        for extra_text in new_paragraphs[len(body_indices):]:
            new_p = _insert_paragraph_after(last_body_para, extra_text, template_info)
            last_body_para = new_p

    _log.info("正文替换完成 | 模板段落数: %d | 新段落数: %d",
              len(body_indices), len(new_paragraphs))


def _insert_paragraph_after(
    ref_para: Paragraph, text: str, template_info: TemplateInfo
) -> Paragraph:
    """在参考段落之后插入新段落。"""
    new_p_element = copy.deepcopy(ref_para._element)
    ref_para._element.addnext(new_p_element)
    new_para = Paragraph(new_p_element, ref_para._parent)

    # 清空并设置文本
    for run in new_para.runs:
        run.text = ""
    if new_para.runs:
        new_para.runs[0].text = text
    else:
        new_para.add_run(text)

    return new_para


def _apply_style(para: Paragraph, template_info: TemplateInfo) -> None:
    """应用模板样式到段落。"""
    if template_info.content_style:
        try:
            para.style = para.part.document.styles[template_info.content_style]
        except KeyError:
            pass
