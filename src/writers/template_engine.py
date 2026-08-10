"""模板引擎 —— 分析 Word 模板结构，提取样式与字段映射。

读取用户的 .docx 模板，分析其结构（标题、字段、段落样式），
为 docx_writer 提供模板信息。
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm

from ..utils.logger import get_logger

_log = get_logger("writer.template_engine")


@dataclass(frozen=True)
class TemplateField:
    """模板字段定义。"""
    label: str          # 字段标签（如"会议类别"）
    value: str          # 模板中的示例值
    style_name: str     # 段落样式名
    paragraph_index: int  # 在文档中的段落位置


@dataclass(frozen=True)
class TemplateInfo:
    """模板解析结果。"""
    title: str                    # 模板标题
    title_style: str              # 标题样式
    doc_number: str               # 文号
    fields: tuple[TemplateField, ...]  # 字段列表
    content_style: str            # 正文段落样式
    content_start_index: int      # 正文起始段落位置
    compiler: str                 # 整理人
    compiler_index: int           # 整理人段落位置
    font_name: str = "宋体"        # 正文字体
    font_size_pt: float = 12.0    # 正文字号


# 已知的模板字段标签前缀
_FIELD_LABELS = [
    "会议类别", "会议议题", "会议时间", "会议地点",
    "主持人", "主 持 人", "参加人员", "纪要内容",
]


def analyze_template(template_path: str | Path) -> TemplateInfo:
    """分析 .docx 模板结构。

    Args:
        template_path: 模板文件路径
    Returns:
        TemplateInfo 模板结构信息
    """
    path = Path(template_path)
    if not path.exists():
        raise FileNotFoundError(f"模板文件不存在: {path}")

    doc = Document(str(path))

    title = ""
    title_style = "Normal"
    doc_number = ""
    fields: list[TemplateField] = []
    content_style = "Normal (Web)"
    content_start_index = -1
    compiler = ""
    compiler_index = -1

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else "Normal"

        # 第一行非空文本是标题
        if not title:
            title = text
            title_style = style
            continue

        # 文号（如 〔2026〕1号）
        if text.startswith("〔") or text.startswith("[") or (text and text[0] in "〔["):
            doc_number = text
            continue

        # 字段行（格式："标签：值"）
        matched = False
        for label in _FIELD_LABELS:
            if label in text:
                # 提取值（冒号后的部分）
                value = text.split("：", 1)[-1].strip() if "：" in text else ""
                fields.append(TemplateField(
                    label=label,
                    value=value,
                    style_name=style,
                    paragraph_index=i,
                ))
                matched = True
                if label == "纪要内容":
                    content_start_index = i
                    content_style = style
                break

        if matched:
            continue

        # 整理人行
        if "整理人" in text:
            compiler = text.split("：", 1)[-1].strip() if "：" in text else ""
            compiler_index = i
            continue

    # 提取字体信息
    font_name = "宋体"
    font_size = 12.0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                font_name = run.font.name
            if run.font.size:
                font_size = run.font.size.pt
            break
        if font_name != "宋体":
            break

    info = TemplateInfo(
        title=title,
        title_style=title_style,
        doc_number=doc_number,
        fields=tuple(fields),
        content_style=content_style,
        content_start_index=content_start_index if content_start_index > 0 else 0,
        compiler=compiler,
        compiler_index=compiler_index if compiler_index > 0 else 0,
        font_name=font_name,
        font_size_pt=font_size,
    )

    _log.info("模板分析完成 | 标题: %s | 字段: %d 个 | 字体: %s %.0fpt",
              info.title, len(info.fields), info.font_name, info.font_size_pt)

    return info
