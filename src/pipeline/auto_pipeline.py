"""全自动写作管线 v2 —— 六层质量飞轮驱动的一键执行引擎。

质量飞轮（Quality Flywheel）：
  L1 范例注入: 生成前，从定稿范文检索黄金段落注入 few-shot
  L2 双盲评审: 生成后，内容官+格式官交叉模型独立打分
  L3 数字指纹: 确定性校验金额/日期/参数防篡改（零 API 成本）
  L4 体裁后处理: Skill 的 post_rules 按体裁执行正则替换
  L5 Self-Refine: 不达标 → 问题清单注入 → 定向修改 → 复审（闭环）
  L6 修改学习: 定稿归档后自动 diff 生成新规则（外部触发）

用户只需提供：输入文件 + 模板（可选）+ 输出目录 + 体裁/子风格（可选）。
"""
from __future__ import annotations

import re
import time
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Optional

from ..config.loader import Config
from ..config.paths import REFINE_DIR
from ..skills import SkillRegistry
from ..llm.multi_llm import MultiLLMClient
from ..processors.understander import understand
from ..processors.generator import generate
from ..processors.refiner import _find_missing_data, _inject_missing_data
from ..processors.post_processor import process as post_process, verify as post_verify
from ..readers.docx_reader import read_docx
from ..readers.transcript_parser import parse_transcript
from ..writers.docx_writer import write_minutes
from ..writers.template_engine import analyze_template
from ..core.models import GeneratedContent
from ..utils.logger import get_logger
from ..workspace.manager import Workspace
from ..quality.exemplar import ExemplarLibrary, build_fewshot_block, detect_genre
from ..quality.fingerprint import audit as fingerprint_audit
from ..quality.self_refine import run_refine_loop

_log = get_logger("pipeline.auto")

ProgressCallback = Callable[[str, int, int], None]
"""回调签名: (步骤描述, 当前步骤号, 总步骤数)"""

# Self-Refine 参数
REFINE_MAX_ROUNDS = 3
REFINE_TARGET_SCORE = 85


@dataclass
class PipelineInput:
    """管线输入参数。"""
    input_files: list[Path] = field(default_factory=list)
    template_files: list[Path] = field(default_factory=list)
    output_dir: Path = field(default_factory=lambda: Path("."))
    sub_genre: str = ""             # 子体裁（如 "会议纪要", "请示件", "函件"）
    sub_style: str = ""             # 子风格（如 "专题会", "调度会"）
    workspace: Optional[Workspace] = None
    style_reference: str = ""       # 预加载的写作风格参考
    revision_guide: str = ""        # 预加载的修改模式指南
    prefer_paid: bool = True
    # 质量飞轮开关
    enable_flywheel: bool = True    # 总开关（False 时退化为 v1 单遍管线）


@dataclass
class PipelineResult:
    """管线输出结果。"""
    success: bool = False
    output_files: list[Path] = field(default_factory=list)
    error: str = ""
    duration: float = 0.0
    quality_issues: list[str] = field(default_factory=list)
    # 飞轮指标
    quality_score: int = 0          # 双盲评审综合分
    refine_rounds: int = 0          # Self-Refine 轮数
    fingerprint_passed: bool = False
    exemplar_count: int = 0         # 注入范例数


def _read_docx_text(path: Path) -> str:
    """读取 docx 纯文本（含文本框）。"""
    try:
        with zipfile.ZipFile(str(path)) as z:
            content = z.read("word/document.xml").decode("utf-8", errors="ignore")
            t_elems = re.findall(r"<w:t[^>]*>([^<]+)</w:t>", content)
            return "".join(t_elems)
    except Exception:
        return read_docx(path)


def _read_any_file(fpath: Path) -> str:
    """按扩展名读取文件文本。"""
    if fpath.suffix == ".docx":
        return _read_docx_text(fpath)
    if fpath.suffix == ".txt":
        return fpath.read_text(encoding="utf-8")
    if fpath.suffix == ".pdf":
        try:
            import fitz
            doc = fitz.open(str(fpath))
            return "".join(p.get_text() for p in doc)
        except Exception:
            return ""
    return _read_docx_text(fpath)


def run_pipeline(
    cfg: Config,
    pipe_input: PipelineInput,
    on_progress: Optional[ProgressCallback] = None,
) -> PipelineResult:
    """执行全自动写作管线（含六层质量飞轮）。"""
    result = PipelineResult()
    start_time = time.time()
    llm = MultiLLMClient(cfg)

    total_steps = 8
    current_step = 0

    def _report(msg: str) -> None:
        nonlocal current_step
        current_step += 1
        if on_progress:
            on_progress(msg, current_step, total_steps)
        _log.info("步骤 %d/%d: %s", current_step, total_steps, msg)

    try:
        # ══ Skill 匹配（L4 规则来源）══
        skill_registry = SkillRegistry()
        matched_skill = skill_registry.match(
            genre=pipe_input.sub_genre,
            sub_style=pipe_input.sub_style,
        )
        skill_prompt = matched_skill.to_prompt_section() if matched_skill else ""
        skill_rules_text = (
            "\n".join(matched_skill.rules) if matched_skill else ""
        )
        if matched_skill:
            _log.info("匹配 Skill: %s (%d 条规则)", matched_skill.name, len(matched_skill.rules))

        # ══ Step 1: 读取输入文件 ══
        _report("读取原始资料")
        combined_text = "\n\n".join(
            filter(None, (_read_any_file(f) for f in pipe_input.input_files))
        ).strip()
        if not combined_text:
            result.error = "原始资料内容为空"
            result.duration = time.time() - start_time
            return result
        _log.info("合并输入: %d字, %d个文件", len(combined_text), len(pipe_input.input_files))

        # ══ Step 2: L1 范例注入准备 + AI 理解 ══
        _report("AI 深度理解内容")

        # L1: 黄金范例检索（体裁 → 相关范例）
        effective_genre = pipe_input.sub_genre or detect_genre(combined_text)
        fewshot_block = ""
        if pipe_input.enable_flywheel:
            try:
                library = ExemplarLibrary(pipe_input.workspace)
                exemplars = library.retrieve(
                    effective_genre, context=combined_text[:2000], k=3,
                )
                fewshot_block = build_fewshot_block(exemplars, effective_genre)
                result.exemplar_count = len(exemplars)
                _log.info("L1 范例注入: %d 段（体裁: %s）", len(exemplars), effective_genre)
            except Exception as e:
                _log.warning("范例库加载失败（跳过 L1）: %s", str(e)[:80])

        transcript = parse_transcript(combined_text)
        analysis = None
        for attempt in range(1, 4):
            try:
                analysis = understand(
                    llm, transcript,
                    prefer_paid=pipe_input.prefer_paid,
                    filename_hint=pipe_input.input_files[0].name if pipe_input.input_files else "",
                )
                topics = len(analysis.topics)
                actions = sum(len(t.action_items) for t in analysis.topics)
                _log.info("理解结果: %d议题, %d行动项", topics, actions)
                if topics >= 2 and actions >= 2:
                    break
            except Exception as e:
                _log.warning("理解尝试 %d 失败: %s", attempt, str(e)[:80])

        if analysis is None:
            result.error = "AI 理解阶段失败"
            result.duration = time.time() - start_time
            return result

        # ══ Step 3: 生成（按模板循环）══
        templates = pipe_input.template_files if pipe_input.template_files else [None]
        total_templates = len(templates)

        for idx, template_path in enumerate(templates):
            tpl_name = template_path.name if template_path else "默认格式"
            if total_templates > 1:
                _report(f"生成文档 ({idx + 1}/{total_templates}): {tpl_name}")
            else:
                _report("AI 正式撰写文档")

            # L1 合并: 范例块 + Skill 块 + 用户风格参考
            style_parts = [p for p in (fewshot_block, skill_prompt, pipe_input.style_reference) if p]
            effective_style = "\n\n".join(style_parts)

            # LLM 生成
            content = generate(
                llm, analysis,
                prefer_paid=pipe_input.prefer_paid,
                style_reference=effective_style,
                revision_guide=pipe_input.revision_guide,
            )

            # 数据注入（防遗漏）
            missing = _find_missing_data(analysis, content)
            if missing:
                injected = _inject_missing_data(list(content.content_paragraphs), missing)
                content = _replace_paragraphs(content, tuple(injected))

            # 确定性后处理（通用规则）
            pp_result = post_process(content)
            content = pp_result.content

            # L4: 体裁专用后处理（Skill post_rules）
            if matched_skill and matched_skill.post_rules:
                content = _apply_skill_post_rules(content, matched_skill.post_rules)
                _log.info("L4 体裁后处理: 应用 %d 条 %s 专用规则",
                          len(matched_skill.post_rules), matched_skill.name)

            # L3 + L2 + L5: 质量飞轮闭环
            if pipe_input.enable_flywheel:
                _report(f"质量飞轮（审计+双盲评审+迭代）: {tpl_name}")
                flywheel_text = _content_to_text(content)
                required_values = _collect_required_values(analysis)
                loop = run_refine_loop(
                    llm,
                    initial_text=flywheel_text,
                    source_text=combined_text,
                    genre=effective_genre,
                    skill_rules=skill_rules_text,
                    max_rounds=REFINE_MAX_ROUNDS,
                    target_score=REFINE_TARGET_SCORE,
                    prefer_paid=pipe_input.prefer_paid,
                    required_values=required_values,
                    on_round=lambda r, s: _log.info("  飞轮 %s", s),
                )
                result.quality_score = loop.scores[-1] if loop.scores else 0
                result.refine_rounds = loop.iterations
                result.fingerprint_passed = loop.audit_passed
                content = _text_to_content(loop.final_text, content)
            else:
                result.fingerprint_passed = True

            # 合规验证
            issues = post_verify(content)
            result.quality_issues.extend(issues)

            # 输出 Word
            pipe_input.output_dir.mkdir(parents=True, exist_ok=True)
            safe_title = content.title.replace("/", "_").replace("\\", "_")[:50]
            output_name = f"{safe_title}.docx"
            output_path = pipe_input.output_dir / output_name
            if template_path:
                template_info = analyze_template(template_path)
                write_minutes(
                    content=content,
                    template_info=template_info,
                    template_path=template_path,
                    output_path=output_path,
                )
            else:
                _write_plain_docx(content, output_path)

            result.output_files.append(output_path)
            _log.info("输出: %s", output_path)

        # ══ Step 4: 验证报告汇总 ══
        _report("合规验证汇总")

        # ══ Step 5: 企业微信推送 ══
        _report("推送成果到企业微信")
        try:
            _push_to_wecom(cfg, result)
        except Exception as e:
            _log.warning("企业微信推送失败: %s", str(e)[:80])

        # ══ Step 6: 完成 ══
        _report("全部完成")
        result.success = True

    except Exception as e:
        result.error = str(e)
        _log.error("管线异常: %s", str(e))

    result.duration = time.time() - start_time
    return result


# ════════════════════════════════════════════════════════
#  内部工具
# ════════════════════════════════════════════════════════

def _collect_required_values(analysis) -> set[str]:
    """从结构化理解结果收集必须出现在成果中的数值（审计白名单）。

    条目格式: "值|单位"（数字对，如 "130|mm"、"0.4|" 无单位留空）
    或完整日期字符串（如 "2026年8月10日"）。
    这些是 AI 理解阶段认定的关键数据，而非转写稿里的口语噪音。
    """
    import re as _re
    num_pat = _re.compile(
        r"(?<![\d月])(\d+(?:\.\d+)?)\s*"
        r"(万元|亿元|mm|cm|km|kN|KN|MPa|Mpa|kV|KV|MW|吨|万t|%|℃|天|日|年|根|米|批|家|人)?"
    )
    date_pat = _re.compile(
        r"\d{4}年\d{1,2}月\d{1,2}日|\d{1,2}月\d{1,2}日"
    )
    values: set[str] = set()

    def _add(text: str) -> None:
        if not text:
            return
        for m in num_pat.finditer(text):
            values.add(f"{m.group(1)}|{m.group(2) or ''}")
        for m in date_pat.finditer(text):
            values.add(m.group(0))

    _add(analysis.meeting_date)
    for topic in analysis.topics:
        for kp in topic.key_data:
            _add(kp.design_value)
            _add(kp.actual_value)
            _add(kp.remark)
        for ai in topic.action_items:
            _add(ai.deadline)

    return values


def _replace_paragraphs(content: GeneratedContent, paragraphs: tuple[str, ...]) -> GeneratedContent:
    """不可变替换段落。"""
    return GeneratedContent(
        title=content.title, doc_number=content.doc_number,
        meeting_type=content.meeting_type, meeting_topic=content.meeting_topic,
        meeting_date=content.meeting_date, meeting_location=content.meeting_location,
        host=content.host, participants=content.participants,
        content_paragraphs=paragraphs,
        compiler=content.compiler, model_used=content.model_used,
    )


def _apply_skill_post_rules(
    content: GeneratedContent,
    post_rules: tuple[tuple[str, str, str], ...],
) -> GeneratedContent:
    """L4: 应用体裁专用后处理规则（正则替换）。"""
    paragraphs = list(content.content_paragraphs)
    applied = 0
    for pattern, replacement, _note in post_rules:
        compiled = re.compile(pattern)
        for i, para in enumerate(paragraphs):
            new_para, n = compiled.subn(replacement, para)
            if n:
                paragraphs[i] = new_para
                applied += n
    if applied:
        _log.info("L4 体裁后处理替换 %d 处", applied)
    return _replace_paragraphs(content, tuple(paragraphs))


def _content_to_text(content: GeneratedContent) -> str:
    """生成全文（标题+正文），供飞轮审计/修改。"""
    parts = [content.title, ""]
    parts.extend(content.content_paragraphs)
    return "\n".join(parts)


def _text_to_content(text: str, original: GeneratedContent) -> GeneratedContent:
    """把飞轮修改后的全文还原为 GeneratedContent（保留元数据）。

    LLM 修改输出常把一句话拆成一行，导致段落碎片化。
    本方法：识别结构性行（标题/称呼/落款/日期/附件），
    其余短行（<60字且非结构）合并进前一段落。
    """
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    if not lines:
        return original

    # 标题识别
    title = original.title
    body_start = 0
    if lines and (
        _title_similarity(lines[0], original.title) > 0.5
        or original.title in lines[0] or lines[0] in original.title
    ):
        title = lines[0]
        body_start = 1

    # 结构性行：保持独立，不合并
    def _is_structural(line: str) -> bool:
        if re.match(r"^(公司领导|各位领导)[：:]", line):
            return True
        if re.match(r"^附件\d*[一二三四五六七八九十]?[：:]", line):
            return True
        if re.match(r"^\d{4}年\d{1,2}月\d{1,2}日$", line):
            return True
        if "事业部" in line and len(line) < 20:
            return True
        return False

    # 话题开头词：以这些词开头的行是新段落，不并入前段
    _topic_opener = re.compile(
        r"^(会议|各单位|各参建|根据|针对|关于|为[确贯落保实]|本次|下一步|"
        r"一、|二、|三、|四、|五、|六、|七、|八、|九、|十、|\d{1,2}[\.、])"
    )

    # 段落重组
    paragraphs: list[str] = []
    for line in lines[body_start:]:
        if _is_structural(line):
            paragraphs.append(line)
        elif not paragraphs:
            paragraphs.append(line)
        elif _topic_opener.match(line):
            paragraphs.append(line)
        elif not paragraphs[-1].endswith(("。", "！", "？", "：", "；")):
            # 前段未以句末标点结尾 → 同段延续
            paragraphs[-1] = paragraphs[-1] + line
        elif len(line) < 60:
            # 短碎片（非话题开头）并入前段
            paragraphs[-1] = paragraphs[-1] + line
        else:
            paragraphs.append(line)

    if not paragraphs:  # 只有标题，正文异常 → 保留原文
        return original
    return GeneratedContent(
        title=title, doc_number=original.doc_number,
        meeting_type=original.meeting_type, meeting_topic=original.meeting_topic,
        meeting_date=original.meeting_date, meeting_location=original.meeting_location,
        host=original.host, participants=original.participants,
        content_paragraphs=tuple(paragraphs),
        compiler=original.compiler, model_used=original.model_used,
    )


def _title_similarity(a: str, b: str) -> float:
    sa, sb = set(a), set(b)
    if not sa or not sb:
        return 0.0
    return len(sa & sb) / len(sa | sb)


def _write_plain_docx(content: GeneratedContent, output_path: Path) -> None:
    """无模板时，生成简单格式 docx。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    h = doc.add_heading(content.title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for para in content.content_paragraphs:
        doc.add_paragraph(para)
    doc.save(str(output_path))


def _push_to_wecom(cfg: Config, result: PipelineResult) -> None:
    """推送成果到企业微信（含质量指标）。"""
    webhook = cfg.get("通知", "企微webhook", "")
    if not webhook or not webhook.startswith("http"):
        _log.info("未配置企业微信 webhook，跳过推送")
        return

    from ..notify.wecom import WeComNotifier
    notifier = WeComNotifier(webhook)
    file_names = ", ".join(f.name for f in result.output_files)
    quality_line = (
        f"\n📊 质量分: {result.quality_score}/100"
        f" | 指纹审计: {'通过' if result.fingerprint_passed else '有问题'}"
        f" | 迭代: {result.refine_rounds}轮"
        if result.quality_score else ""
    )
    notifier.send_text(
        f"✅ Com-Writer 写作完成\n"
        f"📄 生成文件: {file_names}\n"
        f"📂 输出目录: {result.output_files[0].parent if result.output_files else '未知'}"
        f"{quality_line}"
    )
