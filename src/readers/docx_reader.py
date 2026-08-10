"""DOCX 文件读取器 —— 提取段落文本与表格内容。"""
from __future__ import annotations

from pathlib import Path

from docx import Document


def read_docx(file_path: str | Path) -> str:
    """读取 .docx 文件全部文本内容（段落 + 表格）。

    Args:
        file_path: .docx 文件路径
    Returns:
        纯文本内容，段落以 \\n 分隔
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    doc = Document(str(path))
    parts: list[str] = []

    # 段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_docx_paragraphs(file_path: str | Path) -> list[dict]:
    """读取 .docx 文件段落，返回带样式信息的结构。

    Returns:
        [{"text": "...", "style": "Normal", "index": 0}, ...]
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    doc = Document(str(path))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                "text": para.text.strip(),
                "style": para.style.name if para.style else "Normal",
                "index": i,
            })
    return paragraphs


def read_docx_tables(file_path: str | Path) -> list[list[list[str]]]:
    """读取 .docx 文件中的表格数据。

    Returns:
        [[["cell", "cell"], ["cell", "cell"]], ...]  (table -> rows -> cells)
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    doc = Document(str(path))
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return tables
