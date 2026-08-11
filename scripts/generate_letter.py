"""生成发施工单位的正式函件 —— 基于转写稿+模板+参考文件。

流程：
1. 读取转写稿 → 提取函件要点
2. 读取模板 → 获取格式结构
3. 读取参考PDF → 提取政治站位语言
4. LLM生成函件正文
5. 套用模板格式输出Word文档

使用方式：
  编辑下方 CONFIG 区域的路径和名称，或通过 config.ini 配置。
"""
import json
import re
import shutil
import sys
import time
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config.loader import Config
from src.llm.multi_llm import MultiLLMClient
from src.utils.logger import get_logger
from src.config.paths import BASE_DIR, TEMPLATE_DIR, OUTPUT_DIR, RAW_RECORD_DIR, ITERATION_DIR

_log = get_logger("scripts.letter")

# ═══════════════════════════════════════════════════════
# CONFIG — 按实际情况修改以下配置
# ═══════════════════════════════════════════════════════

# 函件标题
LETTER_TITLE_1 = "关于加强对前方支持顺利实现"
LETTER_TITLE_2 = "关键节点目标的函"

# 收件/落款单位（替换为实际公司名称）
RECIPIENT = "XX建设工程局有限公司："
SIGNATURE = "XX规划设计研究院有限公司"
LETTER_DATE = "20XX年X月X日"

# 文件路径
TRANSCRIPT_NAME = ""  # RAW_RECORD_DIR 下最新的 .docx，留空则自动查找
TEMPLATE_NAME = "letter_template.docx"  # TEMPLATE_DIR 下的函件模板
REFERENCE_PDF = ""  # ITERATION_DIR 下的参考PDF（可选）
OUTPUT_PATH = OUTPUT_DIR / "letter_output.docx"

# ═══════════════════════════════════════════════════════


def read_docx_text(path: Path) -> str:
    with zipfile.ZipFile(str(path)) as z:
        content = z.read("word/document.xml").decode("utf-8", errors="ignore")
        t_elems = re.findall(r"<w:t[^>]*>([^<]+)</w:t>", content)
        return "".join(t_elems)


def read_pdf_text(path: Path) -> str:
    import fitz
    doc = fitz.open(str(path))
    return "".join(page.get_text() for page in doc)


_SYSTEM_PROMPT = """你是企业公文写作专家，擅长撰写工程建设领域的正式函件。

═══ 函件写作规则（基于初稿→定稿差异分析，必须遵循）═══

【规则1：格式】标题居中（二号字），抬头顶格（冒号结尾），正文每段首行缩进两字。
落款为公司全称，非部门名。

【规则2：结构精简】函件正文3-4段即可（背景→我方承诺→问题陈述→结尾祝福）。
★★★ 禁止单独写"提高政治站位"的空泛段落——第一段已建立工程重要性，后面不需重复。 ★★★

【规则3：删除过度承诺】★★★ 对外函件中不可写约束己方的过度承诺。
只写己方切实可行的保障措施，不承诺支付方式限制。★★★

【规则4：问题陈述具体化、严重化】★★★ 这是函件说服力的核心。
- 不要泛泛说"资金困难"，要写具体事实
- 使用递进式描述，层层加重
- 必须关联核心目标

【规则5：条件递进后果】问题影响必须用假设条件句强化后果严重性。
格式："如果[问题]得不到明显改善，[中间后果]会严重影响[核心工作]，[核心目标]将难以实现。"

【规则6：要求表述具体可操作】不要写模糊表述。
应给出明确操作指令。

【规则7：语气】"恳请""希望"等商榷语气为主，不卑不亢。用"刚性目标""务必"传达紧迫感。

【规则8：附件】如有支撑文件，在正文后、落款前注明附件名称。

【规则9：转写稿中的现场口语化描述必须转化为正式书面语，但保留生动性和说服力。

【规则10：责任表述】明确"我方已……"和"希望贵司……"的对应关系，体现双方共同努力。

只输出JSON，不要markdown围栏，不要解释。"""

_USER_TEMPLATE = """请根据以下材料，撰写一份正式函件。

【函件标题】{title1}{title2}
【收件单位】{recipient}
【落款单位】{signature}
【日期】{date}

【转写稿要点】（会议讨论的函件内容）：
{transcript}

【参考文件】（工程意义和政治站位表述，仅在第一段适当引用，不单独成段）：
{reference}

【函件模板】（参考格式和行文风格）：
{template}

请输出以下JSON：
{{
  "标题行1": "{title1}",
  "标题行2": "{title2}",
  "抬头": "{recipient}",
  "正文段落": [
    "（第一段：项目背景+工程意义。100-200字）",
    "（第二段：我方承诺和保障措施。80-150字）",
    "（第三段：核心问题段。具体事实→条件递进后果→具体要求。200-350字）",
    "（第四段：结尾祝福。简短，20-40字）"
  ],
  "附件": "",
  "落款": "{signature}",
  "日期": "{date}"
}}

★★★ 关键要求 ★★★
1. 正文仅4段：背景→承诺→问题→祝福。禁止单独写空泛政治段落
2. 问题段必须具体化，使用条件递进后果句式
3. 不写过度承诺
4. 落款为公司全称"""


def _find_transcript() -> Path:
    """查找最新的转写稿。"""
    if TRANSCRIPT_NAME:
        return RAW_RECORD_DIR / TRANSCRIPT_NAME
    docx_files = sorted(RAW_RECORD_DIR.glob("*.docx"))
    if docx_files:
        return docx_files[-1]
    raise FileNotFoundError(f"未找到转写稿: {RAW_RECORD_DIR}")


def main():
    cfg = Config(ROOT / "config.ini")
    llm = MultiLLMClient(cfg)

    # ── 读取所有素材 ──
    _log.info("=" * 60)
    _log.info("读取素材")
    _log.info("=" * 60)

    transcript_path = _find_transcript()
    transcript = read_docx_text(transcript_path)
    _log.info("转写稿: %d字 (%s)", len(transcript), transcript_path.name)

    reference = ""
    if REFERENCE_PDF:
        ref_path = ITERATION_DIR / REFERENCE_PDF
        if ref_path.exists():
            reference = read_pdf_text(ref_path)
            _log.info("参考文件: %d字", len(reference))

    template_path = TEMPLATE_DIR / TEMPLATE_NAME
    template_text = read_docx_text(template_path)
    _log.info("模板: %d字", len(template_text))

    # ── LLM 生成函件 ──
    _log.info("")
    _log.info("=" * 60)
    _log.info("生成函件")
    _log.info("=" * 60)

    start = time.time()
    raw = llm.chat(
        system_prompt=_SYSTEM_PROMPT,
        user_prompt=_USER_TEMPLATE.format(
            title1=LETTER_TITLE_1,
            title2=LETTER_TITLE_2,
            recipient=RECIPIENT,
            signature=SIGNATURE,
            date=LETTER_DATE,
            transcript=transcript[:4000],
            reference=reference[:3000],
            template=template_text,
        ),
        json_mode=True,
        temperature=0.4,
        prefer_paid=True,
        max_tokens=8192,
    )
    elapsed = time.time() - start

    from src.llm.json_utils import extract_json
    data = extract_json(raw)
    _log.info("生成完成 | %.0f秒", elapsed)

    # ── 构建Word文档 ──
    _log.info("")
    _log.info("=" * 60)
    _log.info("套用模板输出Word")
    _log.info("=" * 60)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, OUTPUT_PATH)

    doc = Document(str(OUTPUT_PATH))

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = ""

    paragraphs_data = []
    paragraphs_data.append(("title", data.get("标题行1", "")))
    paragraphs_data.append(("title", data.get("标题行2", "")))
    paragraphs_data.append(("spacer", ""))
    paragraphs_data.append(("body", data.get("抬头", RECIPIENT)))
    for body_text in data.get("正文段落", []):
        if body_text.strip():
            paragraphs_data.append(("body", body_text.strip()))
    attachment = data.get("附件", "").strip()
    if attachment:
        paragraphs_data.append(("body", f"附件：{attachment}"))
    paragraphs_data.append(("spacer", ""))
    paragraphs_data.append(("spacer", ""))
    paragraphs_data.append(("signature", data.get("落款", SIGNATURE)))
    paragraphs_data.append(("date", data.get("日期", LETTER_DATE)))

    doc_paras = doc.paragraphs
    for i, (ptype, text) in enumerate(paragraphs_data):
        if i >= len(doc_paras):
            para = doc.add_paragraph()
        else:
            para = doc_paras[i]

        if ptype == "title":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0] if para.runs else para.add_run()
            run.text = text
            run.font.size = Pt(18)
            run.font.name = "方正小标宋简体"
        elif ptype == "spacer":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif ptype in ("body", "signature"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.runs[0] if para.runs else para.add_run()
            run.text = text
            run.font.size = Pt(16)
        elif ptype == "date":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0] if para.runs else para.add_run()
            run.text = text
            run.font.size = Pt(16)

    while len(doc.paragraphs) > len(paragraphs_data):
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)

    doc.save(str(OUTPUT_PATH))
    _log.info("函件已保存: %s", OUTPUT_PATH)

    # ── 输出预览 ──
    _log.info("")
    _log.info("=" * 60)
    _log.info("【函件预览】")
    _log.info("=" * 60)
    _log.info("标题: %s%s", data.get("标题行1", ""), data.get("标题行2", ""))
    _log.info("抬头: %s", data.get("抬头", ""))
    for i, para_text in enumerate(data.get("正文段落", [])):
        if para_text.strip():
            _log.info("  段落%d (%d字): %s...", i + 1, len(para_text), para_text[:80])
    _log.info("落款: %s %s", data.get("落款", ""), data.get("日期", ""))


if __name__ == "__main__":
    main()
